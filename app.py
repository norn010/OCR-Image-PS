import base64
import difflib
import io
import json
import os
import re
import tempfile
import threading
import time
import uuid
from urllib.parse import parse_qs, unquote_plus, urlparse
from typing import Any, Callable, Optional

import bleach
import markdown
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from flask import Flask, jsonify, render_template, request, send_file
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader, PdfWriter

load_dotenv()


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024  # 30 MB

TYPHOON_OCR_URL = "https://api.opentyphoon.ai/v1/ocr"
OCR_JOBS: dict[str, dict[str, Any]] = {}
OCR_JOBS_LOCK = threading.Lock()
OCR_RESULTS: dict[str, dict[str, Any]] = {}
OCR_RESULTS_LOCK = threading.Lock()
ALLOWED_HTML_TAGS = set(bleach.sanitizer.ALLOWED_TAGS).union(
    {
        "p",
        "br",
        "pre",
        "code",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "table",
        "thead",
        "tbody",
        "tr",
        "th",
        "td",
    }
)


# คำที่ OCR มักอ่านผิดในคอลัมน์รายละเอียด (ใบส่งของ เอส พี บ้านคาร์แคร์)
OCR_DESCRIPTION_CORRECTIONS = [
    ("ล้างอดีต", "ล้างอัดฉีด"),
    ("ฤดูฝน", "ดูดฝุ่น"),
    ("ลงแก้วซากาโยน", "ลงแวกซ์ภายใน"),
    ("ลงแก้วซ์ภายใน", "ลงแวกซ์ภายใน"),
    ("กุณฑนงศักดิ์", "คุณทนงศักดิ์"),
    ("สริพร", "สิริพร"),
    ("S.สริพร", "S.สิริพร"),
]


def correct_ocr_description_text(text: str) -> str:
    """แก้คำที่ OCR อ่านผิดในส่วนรายละเอียด/description"""
    if not text:
        return text
    out = text
    for wrong, right in OCR_DESCRIPTION_CORRECTIONS:
        out = out.replace(wrong, right)
    return out


def ensure_summary_spacing_in_text(text: str) -> str:
    """เว้นบรรทัดก่อนบล็อกสรุป (รวมเงิน/TOTAL, ภาษีมูลค่าเพิ่ม/VAT, รวมสุทธิ/NET TOTAL) ในข้อความที่โชว์"""
    if not text or ("รวมเงิน" not in text and "รวมสุทธิ" not in text):
        return text
    lines = text.splitlines()
    insert_at = -1
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.search(r"รวมเงิน|TOTAL", stripped, re.IGNORECASE) or re.search(r"รวมสุทธิ|NET\s*TOTAL", stripped, re.IGNORECASE):
            insert_at = i
            break
    if insert_at <= 0:
        return text
    blank = ["", ""]
    new_lines = lines[:insert_at] + blank + lines[insert_at:]
    return "\n".join(new_lines)


def render_ocr_html(text: str) -> str:
    html = markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    return bleach.clean(html, tags=ALLOWED_HTML_TAGS, strip=True)


def parse_tables_from_html(html: str) -> list[list[list[str]]]:
    soup = BeautifulSoup(html or "", "html.parser")
    parsed_tables = []
    for table in soup.find_all("table"):
        rows = []
        for tr in table.find_all("tr"):
            cells = tr.find_all(["th", "td"])
            if not cells:
                continue
            row = [" ".join(cell.stripped_strings) for cell in cells]
            rows.append(row)
        if rows:
            parsed_tables.append(rows)
    return parsed_tables


def normalize_header_text(value: str) -> str:
    normalized = re.sub(r"\s+", "", (value or "").lower())
    normalized = re.sub(r"[^0-9a-zก-๙]+", "", normalized)
    return normalized


def header_similarity(left: str, right: str) -> float:
    left_norm = normalize_header_text(left)
    right_norm = normalize_header_text(right)
    if not left_norm or not right_norm:
        return 0.0
    return difflib.SequenceMatcher(None, left_norm, right_norm).ratio()


def is_row_similar_to_header(row: list[str], header: list[str]) -> bool:
    if not row or not header:
        return False
    limit = min(len(row), len(header))
    if limit == 0:
        return False
    matches = 0
    for idx in range(limit):
        if header_similarity(row[idx], header[idx]) >= 0.8:
            matches += 1
    return (matches / limit) >= 0.6


def build_column_mapping(source_header: list[str], target_header: list[str]) -> dict[int, int]:
    mapping: dict[int, int] = {}
    used_targets: set[int] = set()

    for src_idx, src_name in enumerate(source_header):
        best_target = None
        best_score = 0.0
        for tgt_idx, tgt_name in enumerate(target_header):
            if tgt_idx in used_targets:
                continue
            score = header_similarity(src_name, tgt_name)
            if score > best_score:
                best_score = score
                best_target = tgt_idx

        if best_target is not None and best_score >= 0.45:
            mapping[src_idx] = best_target
            used_targets.add(best_target)
        elif src_idx < len(target_header):
            # Fallback by position if header OCR is too noisy.
            mapping[src_idx] = src_idx

    return mapping


def align_row_to_header(
    row: list[str],
    mapping: dict[int, int],
    target_width: int,
) -> list[str]:
    aligned = [""] * target_width
    for src_idx, value in enumerate(row):
        target_idx = mapping.get(src_idx)
        if target_idx is None or target_idx >= target_width:
            continue
        value_clean = (value or "").strip()
        if not value_clean:
            continue
        if aligned[target_idx]:
            aligned[target_idx] = f"{aligned[target_idx]} {value_clean}"
        else:
            aligned[target_idx] = value_clean
    return aligned


def merge_table_rows(tables: list[list[list[str]]]) -> list[list[str]]:
    if not tables:
        return []

    first_table = tables[0]
    if not first_table:
        return []

    primary_header = first_table[0]
    merged_rows: list[list[str]] = [primary_header]
    target_width = len(primary_header)

    first_mapping = {idx: idx for idx in range(target_width)}
    for row in first_table[1:]:
        if is_row_similar_to_header(row, primary_header):
            continue
        merged_rows.append(align_row_to_header(row, first_mapping, target_width))

    for table_rows in tables[1:]:
        if not table_rows:
            continue
        source_header = table_rows[0]
        mapping = build_column_mapping(source_header, primary_header)
        body_rows = table_rows[1:]

        for row in body_rows:
            if is_row_similar_to_header(row, source_header) or is_row_similar_to_header(row, primary_header):
                continue
            merged_rows.append(align_row_to_header(row, mapping, target_width))

    return merged_rows


def build_source_table_payloads(page_htmls: list[str], fallback_html: str) -> list[dict]:
    payloads: list[dict] = []

    if page_htmls:
        for page_number, page_html in enumerate(page_htmls, start=1):
            for table_rows in parse_tables_from_html(page_html):
                if not table_rows:
                    continue
                header = table_rows[0]
                row_order = 0
                rows_with_source = []
                for row in table_rows[1:]:
                    if is_row_similar_to_header(row, header):
                        continue
                    row_order += 1
                    rows_with_source.append((row, f"{page_number}-{row_order}"))
                payloads.append({"header": header, "rows": rows_with_source})
        return payloads

    # Fallback when per-page html is unavailable.
    for table_rows in parse_tables_from_html(fallback_html):
        if not table_rows:
            continue
        header = table_rows[0]
        row_order = 0
        rows_with_source = []
        for row in table_rows[1:]:
            if is_row_similar_to_header(row, header):
                continue
            row_order += 1
            rows_with_source.append((row, f"1-{row_order}"))
        payloads.append({"header": header, "rows": rows_with_source})

    return payloads


def expand_merged_rows_newlines(
    merged_rows: list[list[str]],
    item_col_start: int = 6,
    item_col_end: int = 12,
) -> list[list[str]]:
    """
    Expand data rows where item columns (e.g. 6-11: รหัสสินค้า..จำนวนเงิน) contain newlines
    into one row per line, so Excel/DB get one row per product.
    """
    if not merged_rows or len(merged_rows) < 2:
        return merged_rows
    header = merged_rows[0]
    out: list[list[str]] = [header]
    ncols = len(header)
    for row in merged_rows[1:]:
        padded = list(row) + [""] * (ncols - len(row))
        item_cells = [padded[i] if i < len(padded) else "" for i in range(item_col_start, min(item_col_end, ncols))]
        if not any("\n" in (c or "") for c in item_cells):
            out.append(padded[:ncols])
            continue
        parts = [
            [p.strip() for p in (padded[i] or "").split("\n")]
            for i in range(item_col_start, min(item_col_end, ncols))
        ]
        n = max(len(p) for p in parts) if parts else 1
        for i in range(n):
            new_row = list(padded[:item_col_start])
            for j, part in enumerate(parts):
                new_row.append(part[i] if i < len(part) else "")
            new_row.extend(padded[item_col_end:ncols])
            out.append(new_row[:ncols])
    return out


def merge_table_rows_with_source(table_payloads: list[dict]) -> list[list[str]]:
    if not table_payloads:
        return []

    primary_header = table_payloads[0]["header"]
    target_width = len(primary_header)
    merged_rows: list[list[str]] = [primary_header + ["ที่มาของข้อมูล"]]

    for index, payload in enumerate(table_payloads):
        source_header = payload.get("header", [])
        rows_with_source = payload.get("rows", [])
        if not source_header:
            continue

        if index == 0:
            mapping = {idx: idx for idx in range(target_width)}
        else:
            mapping = build_column_mapping(source_header, primary_header)

        for row, source_ref in rows_with_source:
            if is_row_similar_to_header(row, source_header) or is_row_similar_to_header(row, primary_header):
                continue
            aligned = align_row_to_header(row, mapping, target_width)
            merged_rows.append(aligned + [source_ref])

    return merged_rows


def decode_base64_payload(value: str) -> str:
    if not value:
        return ""
    try:
        return base64.b64decode(value.encode("utf-8")).decode("utf-8")
    except Exception:
        return ""


def decode_base64_json_list(value: str) -> list[str]:
    decoded = decode_base64_payload(value)
    if not decoded:
        return []
    try:
        parsed = json.loads(decoded)
        if isinstance(parsed, list):
            return [str(item) for item in parsed]
    except Exception:
        pass
    return []


def parse_pages_input(pages_raw: str) -> Optional[list[int]]:
    """
    Parse page input to Typhoon pages JSON.
    Supports: "", "all", "1,2,3", "1-5", "1-3,8,10-12"
    Returns list of page numbers or None for all pages.
    """
    raw = (pages_raw or "").strip().lower()
    if not raw or raw == "all":
        return None

    normalized = raw.replace(" ", "")
    tokens = [token for token in normalized.split(",") if token]
    if not tokens:
        return None

    pages_set = set()
    for token in tokens:
        if "-" in token:
            parts = token.split("-")
            if len(parts) != 2 or not parts[0].isdigit() or not parts[1].isdigit():
                raise ValueError("รูปแบบ Pages ไม่ถูกต้อง (ตัวอย่างที่ถูก: 1-39 หรือ 1,2,3)")
            start = int(parts[0])
            end = int(parts[1])
            if start <= 0 or end <= 0:
                raise ValueError("เลขหน้าใน Pages ต้องมากกว่า 0")
            if start > end:
                raise ValueError("ช่วงหน้าใน Pages ต้องเรียงจากน้อยไปมาก เช่น 1-39")
            for page in range(start, end + 1):
                pages_set.add(page)
        else:
            if not token.isdigit():
                raise ValueError("รูปแบบ Pages ไม่ถูกต้อง (ตัวอย่างที่ถูก: 1-39 หรือ 1,2,3)")
            page = int(token)
            if page <= 0:
                raise ValueError("เลขหน้าใน Pages ต้องมากกว่า 0")
            pages_set.add(page)

    if not pages_set:
        return None

    return sorted(pages_set)


def get_pdf_page_count(file_path: str) -> int:
    reader = PdfReader(file_path)
    return len(reader.pages)


ALLOWED_IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg")


def is_image_upload(filename: str, content_type: Optional[str] = None) -> bool:
    if not (filename or "").strip():
        return False
    ext = os.path.splitext(filename.strip())[1].lower()
    if ext in ALLOWED_IMAGE_EXTENSIONS:
        return True
    if content_type and content_type.startswith("image/"):
        return True
    return False


def save_uploaded_file(
    uploaded_bytes: bytes,
    filename: str,
    pdf_password: Optional[str] = None,
) -> tuple[str, str]:
    """
    Save uploaded file to a temp file. Returns (temp_file_path, file_type).
    file_type is "image" or "pdf".
    """
    name = (filename or "").strip() or "document.pdf"
    if is_image_upload(name):
        ext = os.path.splitext(name)[1].lower()
        if ext not in ALLOWED_IMAGE_EXTENSIONS:
            ext = ".png"
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(uploaded_bytes)
            return tmp.name, "image"
    path = save_unlocked_pdf(uploaded_bytes, pdf_password)
    return path, "pdf"


def _parse_native_table_rows(lines: list[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    seen_keys: set[tuple[str, str, str, str, str, str]] = set()
    current_vin = ""
    current_vin_closed = False
    date_range_pattern = r"\d{1,2}/\d{1,2}/\d{4}-\d{1,2}/\d{1,2}/\d{4}"
    money_pattern = r"[0-9,]+\.[0-9]{2}"
    rate_pattern = r"[0-9]+\.[0-9]+"
    combined_row_pattern = re.compile(
        rf"^(?:([A-Z0-9]{{10,}})\s*(\(?Closed\)?)?\s+)?ดอกเบี้ย\s+({date_range_pattern})\s+(\d+)\s+({rate_pattern})\s+({money_pattern})\s+({money_pattern})$",
        flags=re.IGNORECASE,
    )

    def add_row(
        vin: str,
        period: str,
        days: str,
        rate: str,
        principal: str,
        amount: str,
        is_closed: bool,
    ) -> None:
        vin_clean = (vin or "").upper().strip()
        period_clean = (period or "").strip()
        days_clean = (days or "").strip()
        rate_clean = (rate or "").strip()
        principal_clean = (principal or "").strip()
        amount_clean = (amount or "").strip()
        if not vin_clean:
            return
        row_key = (vin_clean, period_clean, days_clean, rate_clean, principal_clean, amount_clean)
        if row_key in seen_keys:
            return
        seen_keys.add(row_key)
        vin_display = f"{vin_clean} (Closed)" if is_closed else vin_clean
        rows.append(
            {
                "vin": vin_display,
                "item": "ดอกเบี้ย",
                "period": period_clean,
                "days": days_clean,
                "rate": rate_clean,
                "principal": principal_clean,
                "amount": amount_clean,
            }
        )

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        # Handle rows where VIN and detail appear on the same line.
        combined_match = combined_row_pattern.match(line)
        if combined_match:
            vin_inline = combined_match.group(1)
            closed_inline = bool(combined_match.group(2))
            vin_value = (vin_inline or current_vin or "").upper().strip()
            if not vin_value:
                continue
            is_closed = closed_inline if vin_inline else current_vin_closed
            add_row(
                vin_value,
                combined_match.group(3),
                combined_match.group(4),
                combined_match.group(5),
                combined_match.group(6),
                combined_match.group(7),
                is_closed,
            )
            current_vin = vin_value
            current_vin_closed = is_closed
            continue

        # Some rows do not include "(Closed)" on the VIN line, so accept VIN-only lines too.
        vin_match = re.match(r"^([A-Z0-9]{10,})(?:\s*(\(?Closed\)?))?", line, flags=re.IGNORECASE)
        if vin_match:
            vin_candidate = vin_match.group(1).upper()
            closed_marker = vin_match.group(2) or ""
            # Guard against accidental header capture: VIN should contain at least one digit.
            if any(ch.isdigit() for ch in vin_candidate):
                current_vin = vin_candidate
                current_vin_closed = "closed" in closed_marker.lower()
                continue

        if not line.startswith("ดอกเบี้ย"):
            continue

        row_pattern = (
            rf"^ดอกเบี้ย\s+({date_range_pattern})\s+(\d+)\s+({rate_pattern})\s+({money_pattern})\s+({money_pattern})$"
        )
        row_match = re.match(row_pattern, line)
        if not row_match:
            continue

        if not current_vin:
            continue

        add_row(
            current_vin,
            row_match.group(1),
            row_match.group(2),
            row_match.group(3),
            row_match.group(4),
            row_match.group(5),
            current_vin_closed,
        )

    # Fallback scan across full page text to catch rows that OCR/text layer split oddly.
    blob = " ".join(lines)
    blob_pattern = re.compile(
        rf"([A-Z0-9]{{10,}})\s*(\(?Closed\)?)?\s*ดอกเบี้ย\s*({date_range_pattern})\s*(\d+)\s*({rate_pattern})\s*({money_pattern})\s*({money_pattern})",
        flags=re.IGNORECASE,
    )
    for match in blob_pattern.finditer(blob):
        add_row(
            match.group(1),
            match.group(3),
            match.group(4),
            match.group(5),
            match.group(6),
            match.group(7),
            bool(match.group(2)),
        )

    return rows


def extract_native_page_content(file_path: str, page_number: int) -> tuple[str, int]:
    reader = PdfReader(file_path)
    if page_number < 1 or page_number > len(reader.pages):
        return "", 0

    text = reader.pages[page_number - 1].extract_text() or ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "", 0

    rows = _parse_native_table_rows(lines)
    if not rows:
        # No parseable table, but keep native text as fallback.
        return "\n".join(lines), 0

    header_lines: list[str] = []
    for line in lines:
        if "เลขตัวถัง/เลขทะเบียน" in line or line.startswith("PRB") or line.startswith("LNN") or line.startswith("LVU"):
            break
        if line in {"-- 1 of 1 --"}:
            continue
        header_lines.append(line)

    md_lines = []
    if header_lines:
        md_lines.append("\n".join(header_lines))
        md_lines.append("")

    md_lines.append("| เลขตัวถัง/เลขทะเบียน | รายการ | ระยะเวลา | วัน | อัตราดอกเบี้ย | ต้นเงินกู้/เงินต้นคงเหลือ | จำนวนเงินที่ต้องชำระ |")
    md_lines.append("|---|---|---|---:|---:|---:|---:|")
    for row in rows:
        md_lines.append(
            f"| {row['vin']} | {row['item']} | {row['period']} | {row['days']} | {row['rate']} | {row['principal']} | {row['amount']} |"
        )

    return "\n".join(md_lines), len(rows)


def init_ocr_job(job_id: str) -> None:
    with OCR_JOBS_LOCK:
        OCR_JOBS[job_id] = {
            "status": "running",
            "message": "กำลังเตรียมไฟล์",
            "current_step": 0,
            "total_steps": 0,
            "current_page_number": 0,
            "page_timings": [],
            "error": "",
            "result": None,
        }


def update_ocr_job(
    job_id: str,
    *,
    status: Optional[str] = None,
    message: Optional[str] = None,
    current_step: Optional[int] = None,
    total_steps: Optional[int] = None,
    current_page_number: Optional[int] = None,
    error: Optional[str] = None,
    result: Optional[dict] = None,
    result_id: Optional[str] = None,
) -> None:
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return
        if status is not None:
            job["status"] = status
        if message is not None:
            job["message"] = message
        if current_step is not None:
            job["current_step"] = current_step
        if total_steps is not None:
            job["total_steps"] = total_steps
        if current_page_number is not None:
            job["current_page_number"] = current_page_number
        if error is not None:
            job["error"] = error
        if result is not None:
            job["result"] = result
        if result_id is not None:
            job["result_id"] = result_id


def append_ocr_job_page_timing(job_id: str, page_number: int, elapsed_seconds: float) -> None:
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return
        timings = job.setdefault("page_timings", [])
        timings.append(
            {
                "page_number": page_number,
                "elapsed_seconds": round(float(elapsed_seconds), 2),
            }
        )


def store_ocr_result(result: dict[str, Any]) -> str:
    result_id = uuid.uuid4().hex
    with OCR_RESULTS_LOCK:
        OCR_RESULTS[result_id] = result
    return result_id


def get_ocr_result(result_id: str) -> Optional[dict[str, Any]]:
    if not result_id:
        return None
    with OCR_RESULTS_LOCK:
        return OCR_RESULTS.get(result_id)


def sanitize_table_name(table_name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]", "_", (table_name or "").strip())
    cleaned = cleaned.strip("_")
    if not cleaned:
        cleaned = "OCR_TTB_WEB"
    return cleaned[:120]


def sanitize_db_name(db_name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]", "", (db_name or "").strip())
    if not cleaned:
        cleaned = "ExcelTtbDB"
    return cleaned[:128]


def get_sql_connection_strings(target_db: str) -> tuple[str, str]:
    """
    Return (master_conn_str, target_db_conn_str) for pyodbc.
    Supports:
    - SQLSERVER_CONNECTION_STRING in SQLAlchemy format:
      mssql+pyodbc://@SERVER/DB?driver=ODBC+Driver+17+for+SQL+Server&trusted_connection=yes
    - fallback to local default instance.
    """
    env_conn = (os.getenv("SQLSERVER_CONNECTION_STRING") or "").strip()
    if env_conn.startswith("mssql+pyodbc://"):
        parsed = urlparse(env_conn)
        server = parsed.netloc.lstrip("@")
        query = parse_qs(parsed.query)
        driver = unquote_plus(query.get("driver", ["ODBC Driver 18 for SQL Server"])[0])
        trusted = query.get("trusted_connection", ["yes"])[0]
        trust_cert = query.get("TrustServerCertificate", ["yes"])[0]
        encrypt = query.get("Encrypt", ["no"])[0]

        base = (
            f"DRIVER={{{driver}}};SERVER={server};"
            f"Trusted_Connection={trusted};TrustServerCertificate={trust_cert};Encrypt={encrypt};"
        )
        return base, base + f"DATABASE={target_db};"

    base = "DRIVER={ODBC Driver 18 for SQL Server};SERVER=.;Trusted_Connection=yes;Encrypt=no;"
    return base, base + f"DATABASE={target_db};"


def normalize_sql_headers(headers: list[str]) -> list[str]:
    normalized: list[str] = []
    seen: dict[str, int] = {}
    for index, header in enumerate(headers, start=1):
        value = re.sub(r"\s+", " ", (header or "").strip())
        if not value:
            value = f"Column_{index}"
        count = seen.get(value, 0) + 1
        seen[value] = count
        if count > 1:
            value = f"{value}_{count}"
        normalized.append(value)
    return normalized


def _clean_cell_for_structured(s: str) -> str:
    if not isinstance(s, str):
        return str(s).strip() if s is not None else ""
    t = str(s).strip()
    t = re.sub(r"รอบกัน", "รอบคัน", t)
    t = re.sub(r"(\d+(?:\.\d*)?)\s+กัน\b", r"\1 คัน", t)
    t = re.sub(r"\s*\(\s*[^)]*บาทถ้วน[^)]*\)\s*", " ", t)
    t = re.sub(r"(60\"?x100FT)\s+(OMODA JAECOO 5)", r"\1 ซันรูฟ \2", t)
    t = re.sub(r"\s+1-\d+\s*$", "", t)
    t = re.sub(r"\s+\d+(?:\.\d*)?\s*คัน\s*$", "", t)
    return re.sub(r"\s+", " ", t).strip()


def extract_ocr_tables_structured(
    html: str,
    fallback_text: str,
    page_htmls: Optional[list[str]] = None,
) -> dict[str, Any]:
    """
    แยกผล OCR เป็น 3 ส่วน: header (หัวเอกสาร), detail (รายการสินค้า), total (สรุปเงิน).
    คืนค่า { "header": {...}, "detail": [[...], ...], "total": {...} }
    """
    table_payloads = build_source_table_payloads(page_htmls or [], html)
    merged_rows = merge_table_rows_with_source(table_payloads)
    txt = (fallback_text or "").strip()

    # Header จากข้อความ OCR
    date_val = inv_val = sales_val = pay_val = due_val = ""
    date_m = re.search(r"วันที่\s*(\d{1,2}/\d{1,2}/\d{4})", txt)
    if date_m:
        date_val = date_m.group(1).strip()
    inv_m = re.search(r"เลขที่\s*([A-Za-z0-9\-]+)", txt)
    if inv_m:
        inv_val = inv_m.group(1).strip()
    emp_m = re.search(r"พนักงานขาย\s*(.+?)(?=\n|กำหนดชำระเงิน|ครบกำหนดวันที่|$)", txt, re.DOTALL)
    if emp_m:
        sales_val = emp_m.group(1).replace("\n", " ").strip()[:200]
    pay_m = re.search(r"กำหนดชำระเงิน\s*(.+?)(?=\n|ครบกำหนดวันที่|$)", txt, re.DOTALL)
    if pay_m:
        pay_val = pay_m.group(1).replace("\n", " ").strip()[:100]
    due_m = re.search(r"ครบกำหนดวันที่\s*(\d{1,2}/\d{1,2}/\d{4})", txt)
    if due_m:
        due_val = due_m.group(1).strip()

    header = {
        "วันที่": date_val,
        "เลขที่": inv_val,
        "พนักงานขาย": sales_val,
        "กำหนดชำระเงิน": pay_val,
        "ครบกำหนดวันที่": due_val,
    }

    # Total (รวมเงิน, VAT, รวมสุทธิ)
    total_val = vat_val = net_val = ""

    def _is_summary_row(row: list) -> bool:
        row_text = " ".join(str(c or "") for c in row).strip()
        first_cell = str((row[0] if row else "") or "").strip()
        if re.match(r"^\d{10,}", first_cell):
            return False
        if first_cell.startswith("คุณ") or re.match(r"^SA\d", first_cell):
            return False
        if not first_cell and re.search(r"คุณ\s+\S+|SA\d{4}-\d+/", row_text):
            return False
        if re.search(r"รวมเงิน\s*/?\s*TOTAL|ภาษีมูลค่าเพิ่ม\s*/?\s*VAT|รวมสุทธิ\s*/?\s*NET\s*TOTAL", row_text, re.IGNORECASE):
            return True
        if re.search(r"(?:แปดร้อย|เก้าร้อย|เจ็ดร้อย|หกร้อย|ห้าร้อย|สี่ร้อย|สามร้อย|สองร้อย|หนึ่งร้อย).*บาท(?:ถ้วน|เพียง|\s*สตางค์)", row_text):
            return True
        if re.search(r"บาท\s*สตางค์", row_text):
            return True
        return False

    def _first_number_in_row(row: list) -> str:
        for cell in row:
            m = re.search(r"([\d,]+\.?\d*)", str(cell or ""))
            if m:
                return m.group(1).strip()
        return ""

    if merged_rows:
        for r in merged_rows[1:]:
            if not _is_summary_row(r):
                continue
            row_text = " ".join(str(c or "") for c in r)
            num = _first_number_in_row(r)
            if not num:
                continue
            if re.search(r"รวมเงิน\s*/?\s*TOTAL", row_text, re.IGNORECASE) and "NET" not in row_text.upper():
                if not total_val:
                    total_val = num
            elif re.search(r"ภาษีมูลค่าเพิ่ม\s*/?\s*VAT|VAT\b", row_text, re.IGNORECASE):
                if not vat_val:
                    vat_val = num
            elif re.search(r"รวมสุทธิ\s*/?\s*NET\s*TOTAL|NET\s*TOTAL", row_text, re.IGNORECASE):
                if not net_val:
                    net_val = num
    if not total_val or not vat_val or not net_val:
        t2 = txt.replace("\n", " ")
        total_m = re.search(r"(?:รวมเงิน|TOTAL)\s*[\s:]*([\d,]+\.?\d*)", t2, re.IGNORECASE)
        vat_m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|VAT)\s*[\s:]*([\d,]+\.?\d*)", t2, re.IGNORECASE)
        net_m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[\s:]*([\d,]+\.?\d*)", t2, re.IGNORECASE)
        if total_m and not total_val:
            total_val = total_m.group(1).strip()
        if vat_m and not vat_val:
            vat_val = vat_m.group(1).strip()
        if net_m and not net_val:
            net_val = net_m.group(1).strip()
    if not total_val or not vat_val or not net_val:
        if not total_val:
            m = re.search(r"(?:รวมเงิน|TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
            if m:
                total_val = m.group(1).strip()
        if not vat_val:
            m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|VAT)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
            if m:
                vat_val = m.group(1).strip()
        if not net_val:
            m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
            if m:
                net_val = m.group(1).strip()

    total = {
        "เลขที่": inv_val,
        "รวมเงิน": total_val,
        "ภาษีมูลค่าเพิ่ม": vat_val,
        "รวมสุทธิ": net_val,
    }

    # Detail: แถวสินค้าจาก merged_rows (กรองแถวสรุป + รวมแถวต่อเนื่อง)
    detail: list[list[str]] = []
    if merged_rows:
        table1_rows = [merged_rows[0]]
        for row in merged_rows[1:]:
            if not _is_summary_row(row):
                table1_rows.append(row)

        def _is_product_code_cell(s: str) -> bool:
            return bool(re.match(r"^\d{10,}", str(s or "").strip()))

        desc_col = 1
        source_ref_col = max(0, len(table1_rows[0]) - 1)
        i = len(table1_rows) - 1
        while i >= 1 and i < len(table1_rows):
            row = table1_rows[i]
            prev = table1_rows[i - 1]
            first_cell = str((row[0] if row else "") or "").strip()
            prev_first = str((prev[0] if prev else "") or "").strip()
            if not _is_product_code_cell(first_cell) and _is_product_code_cell(prev_first):
                extra = " ".join(
                    str(c or "").strip()
                    for ci, c in enumerate(row)
                    if (c and str(c).strip()) and ci != source_ref_col
                )
                if extra:
                    new_prev = list(prev)
                    if desc_col >= len(new_prev):
                        new_prev.extend([""] * (desc_col - len(new_prev) + 1))
                    new_prev[desc_col] = ((new_prev[desc_col] or "").strip() + " " + extra).strip()
                    table1_rows[i - 1] = new_prev
                table1_rows.pop(i)
            else:
                i -= 1

        # คอลัมน์: รหัสสินค้า=0, รายละเอียด=1, จำนวน=2, หน่วยละ=3, %=4, จำนวนเงิน=5
        for row in table1_rows[1:]:
            padded = (list(row) + [""] * 8)[:8]
            doc_no = inv_val
            prod_code = _clean_cell_for_structured(padded[0] if len(padded) > 0 else "")
            desc = _clean_cell_for_structured(padded[1] if len(padded) > 1 else "")
            qty = _clean_cell_for_structured(padded[2] if len(padded) > 2 else "")
            unit_price = _clean_cell_for_structured(padded[3] if len(padded) > 3 else "")
            pct = _clean_cell_for_structured(padded[4] if len(padded) > 4 else "")
            amount = _clean_cell_for_structured(padded[5] if len(padded) > 5 else "")
            detail.append([doc_no, prod_code, desc, qty, unit_price, pct, amount])

    return {"header": header, "detail": detail, "total": total}


def _sql_quote(name: str) -> str:
    return f"[{name.replace(']', ']]')}]"


def parse_preview_sheets_to_structured(sheets_payload: dict[str, Any]) -> dict[str, Any]:
    """
    แปลงข้อมูลจาก Preview Excel (sheets ที่แก้แล้ว) เป็น { header, detail, total }
    โครงสร้าง sheet: แถว 0 = หัวตาราง detail, แถว 1..n = detail, แล้วแถว "รายการ", "รวมเงิน", "VAT", "รวมสุทธิ", แล้ว "วันที่", "เลขที่", ...
    """
    sheets = sheets_payload.get("sheets") or []
    if not sheets:
        return {"header": {}, "detail": [], "total": {}}
    rows = (sheets[0].get("rows") or []) if sheets else []
    if len(rows) < 2:
        return {"header": {}, "detail": [], "total": {}}

    def cell(r: list, c: int) -> str:
        return str((r[c] if c < len(r) else "") or "").strip()

    # หาแถวที่เริ่มตาราง 2 (รายการ)
    table2_start = None
    for i in range(1, len(rows)):
        if "รายการ" in cell(rows[i], 0):
            table2_start = i
            break
    if table2_start is None:
        table2_start = len(rows)

    # total จากแถว รวมเงิน, VAT, รวมสุทธิ
    total_val = vat_val = net_val = ""
    if table2_start + 3 < len(rows):
        total_val = cell(rows[table2_start + 1], 1)
        vat_val = cell(rows[table2_start + 2], 1)
        net_val = cell(rows[table2_start + 3], 1)

    # header จากแถว วันที่, เลขที่, พนักงานขาย, กำหนดชำระเงิน, ครบกำหนดวันที่ (สแกนทั้ง sheet)
    date_val = inv_val = sales_val = pay_val = due_val = ""
    for j in range(len(rows)):
        c0 = cell(rows[j], 0)
        c1 = cell(rows[j], 1)
        if "วันที่" in c0:
            date_val = c1
        elif "เลขที่" in c0 and len(c1) > 1:
            inv_val = c1
        elif "พนักงานขาย" in c0:
            sales_val = c1
        elif "กำหนดชำระเงิน" in c0:
            pay_val = c1
        elif "ครบกำหนดวันที่" in c0:
            due_val = c1
    doc_no = inv_val

    # detail: แถว 1 ถึงก่อน table2_start
    detail: list[list[str]] = []
    for r in range(1, table2_start):
        row = rows[r]
        if not any(cell(row, c) for c in range(6)):
            continue
        prod_code = cell(row, 0)
        desc = cell(row, 1)
        qty = cell(row, 2)
        unit_price = cell(row, 3)
        pct = cell(row, 4)
        amount = cell(row, 5)
        detail.append([doc_no, prod_code, desc, qty, unit_price, pct, amount])

    header = {
        "วันที่": date_val,
        "เลขที่": inv_val,
        "พนักงานขาย": sales_val,
        "กำหนดชำระเงิน": pay_val,
        "ครบกำหนดวันที่": due_val,
    }
    total = {
        "เลขที่": inv_val,
        "รวมเงิน": total_val,
        "ภาษีมูลค่าเพิ่ม": vat_val,
        "รวมสุทธิ": net_val,
    }
    return {"header": header, "detail": detail, "total": total}


def upload_result_to_sql_server(
    result: dict[str, Any], table_name: str, db_name: str = "ExcelTtbDB", edited_sheets_json: Optional[str] = None
) -> dict[str, Any]:
    try:
        import pyodbc  # Local import so app can run without DB dependency.
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "ยังไม่ได้ติดตั้ง pyodbc ใน virtualenv นี้ กรุณารัน: pip install pyodbc"
        ) from exc

    if edited_sheets_json:
        try:
            payload = json.loads(edited_sheets_json)
            data = parse_preview_sheets_to_structured(payload)
            if data is not None:
                header = data.get("header") or {}
                detail = data.get("detail") or []
                total = data.get("total") or {}
            else:
                edited_sheets_json = None
        except (json.JSONDecodeError, KeyError, TypeError):
            edited_sheets_json = None
    if not edited_sheets_json:
        extracted_html = result.get("extracted_html", "")
        extracted_text = result.get("extracted_text", "")
        page_htmls = result.get("page_htmls", []) or []
        data = extract_ocr_tables_structured(extracted_html, extracted_text or "", page_htmls)
        header = data["header"]
        detail = data["detail"]
        total = data["total"]

    safe_base = sanitize_table_name(table_name)
    safe_db = sanitize_db_name(db_name)
    master_conn_str, target_conn_str = get_sql_connection_strings(safe_db)

    conn_master = pyodbc.connect(master_conn_str, autocommit=True)
    cur_master = conn_master.cursor()
    cur_master.execute(f"IF DB_ID('{safe_db}') IS NULL CREATE DATABASE [{safe_db}];")
    conn_master.close()

    conn = pyodbc.connect(target_conn_str)
    cur = conn.cursor()

    # 1. ตาราง header: วันที่, เลขที่, พนักงานขาย, กำหนดชำระเงิน, ครบกำหนดวันที่
    t_header = f"{safe_base}_header"
    cur.execute(f"IF OBJECT_ID('dbo.[{t_header}]', 'U') IS NOT NULL DROP TABLE dbo.[{t_header}];")
    cols_h = ["วันที่", "เลขที่", "พนักงานขาย", "กำหนดชำระเงิน", "ครบกำหนดวันที่"]
    col_defs_h = ", ".join([f"{_sql_quote(c)} NVARCHAR(MAX) NULL" for c in cols_h])
    cur.execute(f"CREATE TABLE dbo.[{t_header}] ({col_defs_h});")
    cur.execute(
        f"INSERT INTO dbo.[{t_header}] ({', '.join(_sql_quote(c) for c in cols_h)}) VALUES (?, ?, ?, ?, ?);",
        [header.get("วันที่", ""), header.get("เลขที่", ""), header.get("พนักงานขาย", ""), header.get("กำหนดชำระเงิน", ""), header.get("ครบกำหนดวันที่", "")],
    )

    # 2. ตาราง detail: เลขที่, รหัสสินค้า, รายละเอียด, จำนวน, หน่วยละ, ส่วนลด, จำนวนเงิน
    t_detail = f"{safe_base}_detail"
    cur.execute(f"IF OBJECT_ID('dbo.[{t_detail}]', 'U') IS NOT NULL DROP TABLE dbo.[{t_detail}];")
    cols_d = ["เลขที่", "รหัสสินค้า", "รายละเอียด", "จำนวน", "หน่วยละ", "ส่วนลด", "จำนวนเงิน"]
    col_defs_d = ", ".join([f"{_sql_quote(c)} NVARCHAR(MAX) NULL" for c in cols_d])
    cur.execute(f"CREATE TABLE dbo.[{t_detail}] ({col_defs_d});")
    if detail:
        insert_d = f"INSERT INTO dbo.[{t_detail}] ({', '.join(_sql_quote(c) for c in cols_d)}) VALUES (?, ?, ?, ?, ?, ?, ?);"
        cur.fast_executemany = True
        cur.executemany(insert_d, [[str(v) if v is not None else "" for v in row] for row in detail])

    # 3. ตาราง total: เลขที่, รวมเงิน, ภาษีมูลค่าเพิ่ม, รวมสุทธิ
    t_total = f"{safe_base}_total"
    cur.execute(f"IF OBJECT_ID('dbo.[{t_total}]', 'U') IS NOT NULL DROP TABLE dbo.[{t_total}];")
    cols_t = ["เลขที่", "รวมเงิน", "ภาษีมูลค่าเพิ่ม", "รวมสุทธิ"]
    col_defs_t = ", ".join([f"{_sql_quote(c)} NVARCHAR(MAX) NULL" for c in cols_t])
    cur.execute(f"CREATE TABLE dbo.[{t_total}] ({col_defs_t});")
    cur.execute(
        f"INSERT INTO dbo.[{t_total}] ({', '.join(_sql_quote(c) for c in cols_t)}) VALUES (?, ?, ?, ?);",
        [total.get("เลขที่", ""), total.get("รวมเงิน", ""), total.get("ภาษีมูลค่าเพิ่ม", ""), total.get("รวมสุทธิ", "")],
    )

    conn.commit()
    cur.execute(f"SELECT COUNT(*) FROM dbo.[{t_detail}];")
    detail_rows = cur.fetchone()[0]
    conn.close()
    return {
        "db_name": safe_db,
        "table_name": f"dbo.{safe_base} (header, detail, total)",
        "rows": 1 + int(detail_rows) + 1,
        "header_table": f"dbo.{t_header}",
        "detail_table": f"dbo.{t_detail}",
        "total_table": f"dbo.{t_total}",
    }


def export_tables_to_docx(
    html: str,
    fallback_text: str,
    page_htmls: Optional[list[str]] = None,
) -> io.BytesIO:
    table_payloads = build_source_table_payloads(page_htmls or [], html)
    merged_rows = merge_table_rows_with_source(table_payloads)

    doc = Document()
    doc.add_heading("OCR Result", level=1)

    if merged_rows:
        max_cols = max(len(row) for row in merged_rows)
        table = doc.add_table(rows=len(merged_rows), cols=max_cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(merged_rows):
            for c_idx in range(max_cols):
                value = row[c_idx] if c_idx < len(row) else ""
                cell = table.cell(r_idx, c_idx)
                cell.text = value
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
    else:
        doc.add_paragraph("No table found in OCR result.")
        if fallback_text.strip():
            doc.add_paragraph(fallback_text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_tables_to_excel(
    html: str,
    fallback_text: str,
    page_htmls: Optional[list[str]] = None,
) -> io.BytesIO:
    table_payloads = build_source_table_payloads(page_htmls or [], html)
    merged_rows = merge_table_rows_with_source(table_payloads)

    wb = Workbook()
    ws = wb.active
    ws.title = "OCR Tables"

    if merged_rows:
        # แก้ OCR ผิด: 'กัน' -> 'คัน' (หน่วย/บริบทรถ) และลบข้อความจำนวนเงินเป็นตัวอักษรที่ปนในรายละเอียด
        def _clean_cell(s: str) -> str:
            if not s or not isinstance(s, str):
                return s
            t = str(s).strip()
            t = re.sub(r"รอบกัน", "รอบคัน", t)
            t = re.sub(r"(\d+(?:\.\d*)?)\s+กัน\b", r"\1 คัน", t)
            t = re.sub(r"\s*\(\s*[^)]*บาทถ้วน[^)]*\)\s*", " ", t)
            # แก้รายการฟิล์มซันรูฟ OMODA JAECOO: เพิ่มคำว่า ซันรูฟ หลัง 60"x100FT ถ้าหายไป
            t = re.sub(r"(60\"?x100FT)\s+(OMODA JAECOO 5)", r"\1 ซันรูฟ \2", t)
            # ลบเลขที่มาของข้อมูลที่ปนในรายละเอียด (เช่น 1-2, 1-3)
            t = re.sub(r"\s+1-\d+\s*$", "", t)
            # ลบจำนวน+หน่วยที่ปนท้ายรายละเอียด (เช่น 1.00 คัน) — ควรอยู่คอลัมน์จำนวน
            t = re.sub(r"\s+\d+(?:\.\d*)?\s*คัน\s*$", "", t)
            return re.sub(r"\s+", " ", t).strip()

        # ไม่เอาแถวสรุป (รวมเงิน/TOTAL, VAT, รวมสุทธิ, จำนวนเงินเป็นตัวอักษร) มาไว้ในตาราง 1 — แสดงแค่ในตาราง 2
        def _is_summary_row(row: list) -> bool:
            row_text = " ".join(str(c or "") for c in row).strip()
            first_cell = str((row[0] if row else "") or "").strip()
            # แถวที่เซลล์แรกเป็นรหัสสินค้า (ตัวเลข 10+ หลัก) เป็นแถวสินค้า ไม่ใช่แถวสรุป
            if re.match(r"^\d{10,}", first_cell):
                return False
            # แถวที่ดูเหมือนส่วนต่อของรายละเอียด (ชื่อคุณ..., หรือว่างแต่มีคุณ/SAเลข) ไม่ถือเป็นแถวสรุป
            if first_cell.startswith("คุณ") or re.match(r"^SA\d", first_cell):
                return False
            if not first_cell and re.search(r"คุณ\s+\S+|SA\d{4}-\d+/", row_text):
                return False
            if re.search(r"รวมเงิน\s*/?\s*TOTAL|ภาษีมูลค่าเพิ่ม\s*/?\s*VAT|รวมสุทธิ\s*/?\s*NET\s*TOTAL", row_text, re.IGNORECASE):
                return True
            # จำนวนเงินเป็นตัวอักษร: ต้องมีคำลงท้ายเช่น บาทถ้วน/บาทเพียง/บาท.*สตางค์
            if re.search(r"(?:แปดร้อย|เก้าร้อย|เจ็ดร้อย|หกร้อย|ห้าร้อย|สี่ร้อย|สามร้อย|สองร้อย|หนึ่งร้อย).*บาท(?:ถ้วน|เพียง|\s*สตางค์)", row_text):
                return True
            if re.search(r"บาท\s*สตางค์", row_text):
                return True
            return False

        table1_rows = [merged_rows[0]]
        for row in merged_rows[1:]:
            if not _is_summary_row(row):
                table1_rows.append(row)

        # รวมแถวที่จริงๆ เป็นส่วนต่อของรายละเอียด (ไม่มีรหัสสินค้าในเซลล์แรก) เข้าไปในแถวสินค้าก่อนหน้า
        def _is_product_code_cell(s: str) -> bool:
            return bool(re.match(r"^\d{10,}", str(s or "").strip()))

        desc_col = 1
        if len(table1_rows[0]) > 1:
            h = str(table1_rows[0][1] or "").upper()
            if "DESCRIPTION" not in h and "รายละเอียด" not in (table1_rows[0][1] or ""):
                for ci, cell in enumerate(table1_rows[0]):
                    if "DESCRIPTION" in str(cell or "").upper() or "รายละเอียด" in str(cell or ""):
                        desc_col = ci
                        break
        # คอลัมน์สุดท้ายคือ ที่มาของข้อมูล (1-1, 1-2, ...) ไม่นำไปต่อท้ายรายละเอียด
        source_ref_col = max(0, len(table1_rows[0]) - 1)
        i = len(table1_rows) - 1
        while i >= 1 and i < len(table1_rows):
            row = table1_rows[i]
            prev = table1_rows[i - 1]
            first_cell = str((row[0] if row else "") or "").strip()
            prev_first = str((prev[0] if prev else "") or "").strip()
            if not _is_product_code_cell(first_cell) and _is_product_code_cell(prev_first):
                extra = " ".join(
                    str(c or "").strip()
                    for ci, c in enumerate(row)
                    if (c and str(c).strip()) and ci != source_ref_col
                )
                if extra:
                    new_prev = list(prev)
                    if desc_col >= len(new_prev):
                        new_prev.extend([""] * (desc_col - len(new_prev) + 1))
                    new_prev[desc_col] = ((new_prev[desc_col] or "").strip() + " " + extra).strip()
                    table1_rows[i - 1] = new_prev
                table1_rows.pop(i)
                # ไม่ลด i เพราะแถวที่เคยอยู่ที่ i+1 เลื่อนมาแทนที่ — ต้องตรวจว่าเป็นแถวต่อเนื่องอีกหรือไม่
            else:
                i -= 1

        max_cols_source = max(len(row) for row in merged_rows)
        table1_cols = min(12, max_cols_source)
        col_widths = [22, 12, 18, 20, 12, 12, 16, 36, 10, 12, 10, 12]
        for c in range(1, table1_cols + 1):
            w = col_widths[c - 1] if c <= len(col_widths) else 12
            ws.column_dimensions[get_column_letter(c)].width = w
        for row in table1_rows:
            row_1 = (row + [""] * table1_cols)[:table1_cols]
            row_1 = [_clean_cell(c) for c in row_1]
            ws.append(row_1)
        header_fill = PatternFill(start_color="D9E2EC", end_color="D9E2EC", fill_type="solid")
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws.freeze_panes = "A2"

        num_table1_rows = len(table1_rows)
        # ช่วงว่าง 1 แถว ระหว่างตาราง 1 กับตาราง 2
        gap_row = num_table1_rows + 1
        ws.row_dimensions[gap_row].height = 18
        for c in range(1, table1_cols + 1):
            ws.cell(row=gap_row, column=c, value="")

        # ตาราง 2: สรุป (รวมเงิน, VAT, รวมสุทธิ) — ดึงจากแถวสรุปใน merged_rows ก่อน แล้วค่อยจากข้อความ OCR
        def _first_number_in_row(row: list) -> str:
            for cell in row:
                m = re.search(r"([\d,]+\.?\d*)", str(cell or ""))
                if m:
                    return m.group(1).strip()
            return ""

        total_val = vat_val = net_val = ""
        for r in merged_rows[1:]:
            if not _is_summary_row(r):
                continue
            row_text = " ".join(str(c or "") for c in r)
            num = _first_number_in_row(r)
            if not num:
                continue
            if re.search(r"รวมเงิน\s*/?\s*TOTAL", row_text, re.IGNORECASE) and "NET" not in row_text.upper():
                if not total_val:
                    total_val = num
            elif re.search(r"ภาษีมูลค่าเพิ่ม\s*/?\s*VAT|VAT\b", row_text, re.IGNORECASE):
                if not vat_val:
                    vat_val = num
            elif re.search(r"รวมสุทธิ\s*/?\s*NET\s*TOTAL|NET\s*TOTAL", row_text, re.IGNORECASE):
                if not net_val:
                    net_val = num
        if not total_val or not vat_val or not net_val:
            txt = (fallback_text or "").replace("\n", " ")
            total_m = re.search(r"(?:รวมเงิน|TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE)
            vat_m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|VAT)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE)
            net_m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE)
            if total_m and not total_val:
                total_val = total_m.group(1).strip()
            if vat_m and not vat_val:
                vat_val = vat_m.group(1).strip()
            if net_m and not net_val:
                net_val = net_m.group(1).strip()
        # fallback: ถ้าข้อความมีหลายบรรทัด ให้ลอง match แบบข้ามบรรทัด (หาตัวเลขที่อยู่หลัง label ไม่เกิน ~80 ตัวอักษร)
        if not total_val or not vat_val or not net_val:
            txt = fallback_text or ""
            if not total_val:
                total_m = re.search(r"(?:รวมเงิน|TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
                if total_m:
                    total_val = total_m.group(1).strip()
            if not vat_val:
                vat_m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|VAT)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
                if vat_m:
                    vat_val = vat_m.group(1).strip()
            if not net_val:
                net_m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[\s:]*([\d,]+\.?\d*)", txt, re.IGNORECASE | re.DOTALL)
                if net_m:
                    net_val = net_m.group(1).strip()
        table2_start_row = gap_row + 1
        ws.cell(row=table2_start_row, column=1, value="รายการ")
        ws.cell(row=table2_start_row, column=2, value="จำนวน")
        ws.cell(row=table2_start_row + 1, column=1, value="รวมเงิน/TOTAL")
        ws.cell(row=table2_start_row + 1, column=2, value=total_val)
        ws.cell(row=table2_start_row + 2, column=1, value="ภาษีมูลค่าเพิ่ม/VAT")
        ws.cell(row=table2_start_row + 2, column=2, value=vat_val)
        ws.cell(row=table2_start_row + 3, column=1, value="รวมสุทธิ/NET TOTAL")
        ws.cell(row=table2_start_row + 3, column=2, value=net_val)
        for r in range(table2_start_row, table2_start_row + 4):
            for col in (1, 2):
                if r == table2_start_row:
                    ws.cell(row=r, column=col).font = Font(bold=True)
        ws.column_dimensions["A"].width = 28
        ws.column_dimensions["B"].width = 14

        # ตาราง 3: ข้อมูลหัวเอกสาร — ดึงจากข้อความ OCR เท่านั้น (ไม่ใช้จากแถวสินค้าในตาราง 1)
        date_val = inv_val = sales_val = pay_val = due_val = ""
        txt = fallback_text or ""
        date_m = re.search(r"วันที่\s*(\d{1,2}/\d{1,2}/\d{4})", txt)
        if date_m:
            date_val = date_m.group(1).strip()
        inv_m = re.search(r"เลขที่\s*([A-Za-z0-9\-]+)", txt)
        if inv_m:
            inv_val = inv_m.group(1).strip()
        emp_m = re.search(r"พนักงานขาย\s*(.+?)(?=\n|กำหนดชำระเงิน|ครบกำหนดวันที่|$)", txt, re.DOTALL)
        if emp_m:
            sales_val = emp_m.group(1).replace("\n", " ").strip()[:200]
        pay_m = re.search(r"กำหนดชำระเงิน\s*(.+?)(?=\n|ครบกำหนดวันที่|$)", txt, re.DOTALL)
        if pay_m:
            pay_val = pay_m.group(1).replace("\n", " ").strip()[:100]
        due_m = re.search(r"ครบกำหนดวันที่\s*(\d{1,2}/\d{1,2}/\d{4})", txt)
        if due_m:
            due_val = due_m.group(1).strip()
        gap3_row = table2_start_row + 4
        ws.row_dimensions[gap3_row].height = 18
        ws.cell(row=gap3_row, column=1, value="")
        ws.cell(row=gap3_row, column=2, value="")
        table3_start = gap3_row + 1
        ws.cell(row=table3_start, column=1, value="วันที่")
        ws.cell(row=table3_start, column=2, value=date_val)
        ws.cell(row=table3_start + 1, column=1, value="เลขที่")
        ws.cell(row=table3_start + 1, column=2, value=inv_val)
        ws.cell(row=table3_start + 2, column=1, value="พนักงานขาย")
        ws.cell(row=table3_start + 2, column=2, value=sales_val)
        ws.cell(row=table3_start + 3, column=1, value="กำหนดชำระเงิน")
        ws.cell(row=table3_start + 3, column=2, value=pay_val)
        ws.cell(row=table3_start + 4, column=1, value="ครบกำหนดวันที่")
        ws.cell(row=table3_start + 4, column=2, value=due_val)
    else:
        # กรณีไม่มีตาราง (เช่น ผล OCR เป็นข้อความล้วน): เขียนข้อความทีละบรรทัด แล้วเว้น 2 แถว แล้วเขียนสรุป
        if fallback_text.strip():
            for line in fallback_text.splitlines():
                ws.append([line.strip()])
        else:
            ws.append(["No table found in OCR result."])
        last_row = ws.max_row
        # เว้นแถวว่าง 2 แถว (ตั้งความสูงให้มองเห็น)
        for i in range(1, 3):
            r = last_row + i
            ws.row_dimensions[r].height = 18
            ws.cell(row=r, column=1, value="")
        # พยายามดึง รวมเงิน, VAT, รวมสุทธิ จากข้อความแล้วเขียนต่อ
        total_m = re.search(r"(?:รวมเงิน|TOTAL)\s*[\s:]*([\d,]+\.?\d*)", fallback_text or "", re.IGNORECASE)
        vat_m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|VAT)\s*[\s:]*([\d,]+\.?\d*)", fallback_text or "", re.IGNORECASE)
        net_m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[\s:]*([\d,]+\.?\d*)", fallback_text or "", re.IGNORECASE)
        summary_start = last_row + 3
        if total_m:
            ws.cell(row=summary_start, column=1, value="รวมเงิน/TOTAL")
            ws.cell(row=summary_start, column=2, value=total_m.group(1).strip())
        if vat_m:
            ws.cell(row=summary_start + 1, column=1, value="ภาษีมูลค่าเพิ่ม/VAT")
            ws.cell(row=summary_start + 1, column=2, value=vat_m.group(1).strip())
        if net_m:
            ws.cell(row=summary_start + 2, column=1, value="รวมสุทธิ/NET TOTAL")
            ws.cell(row=summary_start + 2, column=2, value=net_m.group(1).strip())

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def save_unlocked_pdf(uploaded_bytes: bytes, password: Optional[str]) -> str:
    """Return path to a temporary unlocked PDF file."""
    reader = PdfReader(io.BytesIO(uploaded_bytes))

    if reader.is_encrypted:
        if not password:
            raise ValueError("ไฟล์ PDF นี้ถูกล็อก กรุณาใส่รหัสผ่าน PDF")
        if reader.decrypt(password) == 0:
            raise ValueError("รหัสผ่าน PDF ไม่ถูกต้อง")

    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        writer.write(tmp_file)
        return tmp_file.name


def call_typhoon_ocr_single_request(
    file_path: str,
    api_key: str,
    model: str,
    task_type: str,
    max_tokens: int,
    temperature: float,
    top_p: float,
    repetition_penalty: float,
    pages_json: Optional[str] = None,
) -> tuple[str, list[str]]:
    data = {
        "model": model,
        "task_type": task_type,
        "max_tokens": str(max_tokens),
        "temperature": str(temperature),
        "top_p": str(top_p),
        "repetition_penalty": str(repetition_penalty),
    }
    if pages_json:
        data["pages"] = pages_json

    headers = {"Authorization": f"Bearer {api_key}"}

    with open(file_path, "rb") as file:
        response = requests.post(
            TYPHOON_OCR_URL,
            files={"file": file},
            data=data,
            headers=headers,
            timeout=180,
        )

    if response.status_code != 200:
        raise RuntimeError(f"Typhoon API error {response.status_code}: {response.text}")

    result = response.json()
    extracted_texts = []
    per_page_texts: list[str] = []
    for page_result in result.get("results", []):
        if page_result.get("success") and page_result.get("message"):
            content = page_result["message"]["choices"][0]["message"]["content"]
            try:
                parsed_content = json.loads(content)
                text = parsed_content.get("natural_text", content)
            except json.JSONDecodeError:
                text = content
            extracted_texts.append(text)
            per_page_texts.append(text)
        elif not page_result.get("success"):
            error_msg = page_result.get("error", "Unknown error")
            extracted_texts.append(f"[ERROR] {error_msg}")
            per_page_texts.append(f"[ERROR] {error_msg}")

    return "\n\n".join(extracted_texts), per_page_texts


def call_typhoon_ocr(
    file_path: str,
    api_key: str,
    model: str,
    task_type: str,
    max_tokens: int,
    temperature: float,
    top_p: float,
    repetition_penalty: float,
    pages: Optional[list[int]] = None,
    progress_callback: Optional[Callable[[int, int, int], None]] = None,
    page_done_callback: Optional[Callable[[int, int, int, float], None]] = None,
    use_native_extraction: bool = True,
) -> tuple[str, list[str], list[dict[str, Any]]]:
    """
    OCR strategy:
    - If multiple pages are requested, call API page-by-page for better completeness.
    - If one page (or unknown page list), do a single call.
    - use_native_extraction=False for image uploads (no PDF text layer).
    """
    if pages and len(pages) > 1:
        merged_texts: list[str] = []
        per_page_texts: list[str] = []
        page_timings: list[dict[str, Any]] = []
        total_pages = len(pages)

        for page_index, page_number in enumerate(pages, start=1):
            if progress_callback:
                progress_callback(page_index, total_pages, page_number)
            page_started = time.perf_counter()

            if use_native_extraction:
                native_text, native_row_count = extract_native_page_content(file_path, page_number)
            else:
                native_text, native_row_count = "", 0
            if native_text and native_row_count > 0:
                merged_texts.append(native_text)
                per_page_texts.append(native_text)
                elapsed = round(time.perf_counter() - page_started, 2)
                page_timings.append({"page_number": page_number, "elapsed_seconds": elapsed})
                if page_done_callback:
                    page_done_callback(page_index, total_pages, page_number, elapsed)
                continue

            page_payload = json.dumps([page_number])
            last_error: Optional[Exception] = None

            # Retry each page a few times to reduce transient misses/timeouts.
            for _ in range(3):
                try:
                    page_joined, page_results = call_typhoon_ocr_single_request(
                        file_path=file_path,
                        api_key=api_key,
                        model=model,
                        task_type=task_type,
                        max_tokens=max_tokens,
                        temperature=temperature,
                        top_p=top_p,
                        repetition_penalty=repetition_penalty,
                        pages_json=page_payload,
                    )
                    if page_results:
                        page_text = page_results[0]
                    else:
                        page_text = page_joined
                    merged_texts.append(page_text)
                    per_page_texts.append(page_text)
                    last_error = None
                    break
                except Exception as exc:
                    last_error = exc

            if last_error is not None:
                raise RuntimeError(f"OCR failed on page {page_number}: {last_error}") from last_error

            elapsed = round(time.perf_counter() - page_started, 2)
            page_timings.append({"page_number": page_number, "elapsed_seconds": elapsed})
            if page_done_callback:
                page_done_callback(page_index, total_pages, page_number, elapsed)

        return "\n\n".join(merged_texts), per_page_texts, page_timings

    if progress_callback:
        single_page_number = pages[0] if pages else 1
        progress_callback(1, 1, single_page_number)

    started = time.perf_counter()
    if pages and len(pages) == 1 and use_native_extraction:
        native_text, native_row_count = extract_native_page_content(file_path, pages[0])
        if native_text and native_row_count > 0:
            elapsed = round(time.perf_counter() - started, 2)
            page_timings = [{"page_number": pages[0], "elapsed_seconds": elapsed}]
            if page_done_callback:
                page_done_callback(1, 1, pages[0], elapsed)
            return native_text, [native_text], page_timings

    pages_json = json.dumps(pages) if pages else None
    merged_text, per_page_texts = call_typhoon_ocr_single_request(
        file_path=file_path,
        api_key=api_key,
        model=model,
        task_type=task_type,
        max_tokens=max_tokens,
        temperature=temperature,
        top_p=top_p,
        repetition_penalty=repetition_penalty,
        pages_json=pages_json,
    )

    elapsed = round(time.perf_counter() - started, 2)
    page_number = pages[0] if pages else 1
    page_timings = [{"page_number": page_number, "elapsed_seconds": elapsed}]
    if page_done_callback:
        page_done_callback(1, 1, page_number, elapsed)
    return merged_text, per_page_texts, page_timings


# SP Baan Car Care (เอส พี บ้านคาร์แคร์) — ใบส่งของ/ใบกำกับภาษี
SP_TABLE_HEADER = (
    "บริษัท|วันที่|เลขที่|พนักงานขาย|กำหนดชำระเงิน|ครบกำหนดวันที่|"
    "รหัสสินค้า/PROD.CODE|รายละเอียด/DESCRIPTION|จำนวน/QUANTITY|"
    "หน่วยละ/UNIT/PRICE|ส่วนลด%|จำนวนเงิน/AMOUNT|"
    "รวมเงิน/total|ภาษีมูลค่ารวม/vat|รวมสุทธิ/net total"
)


def _extract_sp_header(text: str) -> dict[str, str]:
    """Extract header fields from SP Baan Car Care OCR text."""
    out: dict[str, str] = {
        "บริษัท": "",
        "วันที่": "",
        "เลขที่": "",
        "พนักงานขาย": "",
        "กำหนดชำระเงิน": "",
        "ครบกำหนดวันที่": "",
        "รวมเงิน/total": "",
        "ภาษีมูลค่ารวม/vat": "",
        "รวมสุทธิ/net total": "",
    }
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]

    for i, line in enumerate(lines):
        if re.search(r"บริษัท\s+.+คาร์แคร์|เอส\s*พี\s*บ้านคาร์แคร์", line):
            out["บริษัท"] = line.replace("บริษัท", "").strip()
            if not out["บริษัท"] and i > 0:
                out["บริษัท"] = "บริษัท เอส พี บ้านคาร์แคร์ จำกัด (สำนักงานใหญ่)"
            if not out["บริษัท"]:
                out["บริษัท"] = line.strip()
            break

    date_m = re.search(r"วันที่\s*(\d{1,2}/\d{1,2}/\d{4})", text or "")
    if date_m:
        out["วันที่"] = date_m.group(1)

    inv_m = re.search(r"เลขที่\s*([A-Za-z0-9\-]+)", text or "")
    if inv_m:
        out["เลขที่"] = inv_m.group(1).strip()

    emp_m = re.search(r"พนักงานขาย\s*(.+?)(?=\n|กำหนดชำระเงิน|ครบกำหนด|$)", text or "", re.DOTALL)
    if emp_m:
        out["พนักงานขาย"] = emp_m.group(1).replace("\n", " ").strip()[:200]

    pay_m = re.search(r"กำหนดชำระเงิน\s*(.+?)(?=\n|ครบกำหนดวันที่|$)", text or "", re.DOTALL)
    if pay_m:
        out["กำหนดชำระเงิน"] = pay_m.group(1).replace("\n", " ").strip()[:100]

    due_m = re.search(r"ครบกำหนดวันที่\s*(\d{1,2}/\d{1,2}/\d{4})", text or "")
    if due_m:
        out["ครบกำหนดวันที่"] = due_m.group(1)

    # Summary: รวมเงิน/TOTAL, ภาษีมูลค่าเพิ่ม/VAT, รวมสุทธิ/NET TOTAL
    total_m = re.search(r"(?:รวมเงิน|TOTAL)\s*[:\s]*([\d,]+\.?\d*)", text or "", re.IGNORECASE)
    if total_m:
        out["รวมเงิน/total"] = total_m.group(1).strip()
    vat_m = re.search(r"(?:ภาษีมูลค่าเพิ่ม|ภาษีมูลค่ารวม|VAT)\s*[:\s]*([\d,]+\.?\d*)", text or "", re.IGNORECASE)
    if vat_m:
        out["ภาษีมูลค่ารวม/vat"] = vat_m.group(1).strip()
    net_m = re.search(r"(?:รวมสุทธิ|NET\s*TOTAL)\s*[:\s]*([\d,]+\.?\d*)", text or "", re.IGNORECASE)
    if net_m:
        out["รวมสุทธิ/net total"] = net_m.group(1).strip()

    return out


def _is_product_line_item(code: str, desc: str) -> bool:
    """
    กรองเฉพาะแถวที่เป็นรายการสินค้าจริง ไม่เอาแถวที่เป็นเลขประจำตัวผู้เสียภาษี/รหัสลูกค้า/หัวตาราง
    รหัสสินค้าจริงมักขึ้นต้นด้วย 89 (เช่น 8901000000100) เลขประจำตัวผู้เสียภาษีมักขึ้นต้น 0 (เช่น 0305538001196)
    """
    code_clean = (code or "").strip().replace(" ", "")
    desc_clean = (desc or "").strip()
    if not code_clean or len(code_clean) < 10:
        return False
    if re.match(r"^0\d{12}$", code_clean):
        return False
    skip_desc = ("รหัสลูกค้า", "นามลูกค้า", "ที่อยู่", "ใบส่งของ/ใบกำกับภาษี", "DELIVERY ORDER/TAX INVOICE")
    if any(s in desc_clean for s in skip_desc) and len(desc_clean) < 80:
        return False
    return True


def _extract_sp_line_items(text: str) -> list[tuple[str, str, str, str, str, str]]:
    """
    Extract line items: (รหัสสินค้า, รายละเอียด, จำนวน, หน่วยละ, ส่วนลด%, จำนวนเงิน).
    Tries markdown table first (if Typhoon returns one), then regex on plain lines.
    กรองเฉพาะรายการสินค้าจริง (ไม่เอาเลขประจำตัวผู้เสียภาษี/รหัสลูกค้า)
    """
    items: list[tuple[str, str, str, str, str, str]] = []
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]

    # Try parsing markdown table: | col1 | col2 | ...
    table_rows: list[list[str]] = []
    for line in lines:
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not cells or all(re.match(r"^\-+$", c) for c in cells):
                continue
            table_rows.append(cells)
    if table_rows and len(table_rows) >= 2:
        header_row = table_rows[0]
        # Map header to: รหัสสินค้า, รายละเอียด, จำนวน, หน่วยละ, ส่วนลด%, จำนวนเงิน
        idx_code = idx_desc = idx_qty = idx_unit = idx_disc = idx_amt = -1
        for i, h in enumerate(header_row):
            n = normalize_header_text(h)
            if "รหัสสินค้า" in n or "prod" in n or "code" in n:
                idx_code = i
            elif "รายละเอียด" in n or "description" in n:
                idx_desc = i
            elif "จำนวน" in n and "เงิน" not in n or "quantity" in n:
                idx_qty = i
            elif "หน่วยละ" in n or "unit" in n or "price" in n:
                idx_unit = i
            elif "ส่วนลด" in n or "discount" in n:
                idx_disc = i
            elif "จำนวนเงิน" in n or "amount" in n:
                idx_amt = i
        if idx_code >= 0 or idx_desc >= 0:
            def get_cell(row: list[str], idx: int) -> str:
                return (row[idx] if 0 <= idx < len(row) else "").strip()
            for row in table_rows[1:]:
                if is_row_similar_to_header(row, header_row):
                    continue
                code = get_cell(row, idx_code)
                desc = get_cell(row, idx_desc)
                if not _is_product_line_item(code, desc):
                    continue
                items.append((
                    code,
                    desc,
                    get_cell(row, idx_qty),
                    get_cell(row, idx_unit),
                    get_cell(row, idx_disc),
                    get_cell(row, idx_amt),
                ))
            if items:
                return items
    # Fallback: regex on plain lines (product code 10+ digits, then description, then quantity/unit/discount/amount)
    num = r"[\d,]+\.?\d*"
    for line in lines:
        if not line or len(line) < 5:
            continue
        # Match 13-digit code, then description (until we hit numbers), then quantity, unit price, discount?, amount
        m = re.match(
            r"^(\d{10,})\s+(.+?)\s+(" + num + r")\s*(?:คัน|หน่วย|ชิ้น)?\s*(" + num + r")?\s*(" + num + r")?\s*(" + num + r")?$",
            line,
        )
        if m:
            code, desc, qty = m.group(1), m.group(2).strip(), m.group(3)
            unit = (m.group(4) or "").strip()
            disc = (m.group(5) or "").strip()
            amt = (m.group(6) or "").strip()
            if _is_product_line_item(code, desc):
                items.append((code, desc, qty, unit, disc, amt))
            continue
        # Fallback: line with at least one long numeric code and some numbers
        parts = re.split(r"\s{2,}|\t", line)
        if len(parts) >= 4 and re.match(r"^\d{10,}$", (parts[0] or "").replace(" ", "")):
            code = (parts[0] or "").strip()
            desc = (parts[1] or "").strip() if len(parts) > 1 else ""
            if not _is_product_line_item(code, desc):
                continue
            rest = [p.strip() for p in parts[2:] if p.strip()]
            qty = rest[0] if len(rest) > 0 else ""
            unit = rest[1] if len(rest) > 1 else ""
            disc = rest[2] if len(rest) > 2 else ""
            amt = rest[3] if len(rest) > 3 else ""
            items.append((code, desc, qty, unit, disc, amt))

    # Alternative: scan for lines that look like "8901000000100 ... 280.00" (code + description + numbers)
    if not items:
        blob = " ".join(lines)
        # Pattern: 13 digits, then non-digit text, then several decimal numbers
        for m in re.finditer(
            r"(\d{10,})\s+([^\d]+?)\s+(" + num + r")\s*(?:คัน|หน่วย|ชิ้น)?\s*(" + num + r")?\s*(" + num + r")?\s*(" + num + r")?",
            blob,
        ):
            code = m.group(1)
            desc = m.group(2).strip()
            qty = m.group(3)
            unit = (m.group(4) or "").strip()
            disc = (m.group(5) or "").strip()
            amt = (m.group(6) or "").strip()
            if desc and (unit or amt or qty) and _is_product_line_item(code, desc):
                items.append((code, desc, qty, unit, disc, amt))

    return items


def _expand_sp_line_items(
    items: list[tuple[str, str, str, str, str, str]]
) -> list[tuple[str, str, str, str, str, str]]:
    """
    Expand any line item whose fields contain newlines into multiple items (one per line).
    So one row with '8901...\n8905...' and 'desc1\ndesc2' becomes two rows.
    """
    expanded: list[tuple[str, str, str, str, str, str]] = []
    for (code, desc, qty, unit, disc, amt) in items:
        if "\n" not in code and "\n" not in desc and "\n" not in qty and "\n" not in unit and "\n" not in disc and "\n" not in amt:
            expanded.append((code, desc, qty, unit, disc, amt))
            continue
        parts_code = [p.strip() for p in code.split("\n")]
        parts_desc = [p.strip() for p in desc.split("\n")]
        parts_qty = [p.strip() for p in qty.split("\n")]
        parts_unit = [p.strip() for p in unit.split("\n")]
        parts_disc = [p.strip() for p in disc.split("\n")]
        parts_amt = [p.strip() for p in amt.split("\n")]
        n = max(
            len(parts_code), len(parts_desc), len(parts_qty),
            len(parts_unit), len(parts_disc), len(parts_amt),
            1,
        )
        for i in range(n):
            expanded.append((
                parts_code[i] if i < len(parts_code) else "",
                parts_desc[i] if i < len(parts_desc) else "",
                parts_qty[i] if i < len(parts_qty) else "",
                parts_unit[i] if i < len(parts_unit) else "",
                parts_disc[i] if i < len(parts_disc) else "",
                parts_amt[i] if i < len(parts_amt) else "",
            ))
    return expanded


def parse_sp_baan_care_to_markdown(ocr_text: str) -> str:
    """
    Parse OCR text from SP Baan Car Care delivery order/invoice into a markdown table
    with 15 columns (including รวมเงิน/total, ภาษีมูลค่ารวม/vat, รวมสุทธิ/net total). Returns empty string if no line items found.
    """
    if not (ocr_text or "").strip():
        return ""
    header = _extract_sp_header(ocr_text)
    items = _extract_sp_line_items(ocr_text)
    if not items:
        return ""
    company = header.get("บริษัท") or "บริษัท เอส พี บ้านคาร์แคร์ จำกัด (สำนักงานใหญ่)"
    date = header.get("วันที่") or ""
    inv = header.get("เลขที่") or ""
    sales = header.get("พนักงานขาย") or ""
    pay_terms = header.get("กำหนดชำระเงิน") or ""
    due = header.get("ครบกำหนดวันที่") or ""
    total = header.get("รวมเงิน/total") or ""
    vat = header.get("ภาษีมูลค่ารวม/vat") or ""
    net_total = header.get("รวมสุทธิ/net total") or ""

    # ใส่หลายรายการในแถวเดียวกัน คั่นด้วย \n (ตามรูปที่ user ต้องการ)
    codes = "\n".join(it[0] for it in items)
    descs = "\n".join(it[1] for it in items)
    qtys = "\n".join(it[2] for it in items)
    units = "\n".join(it[3] for it in items)
    discs = "\n".join(it[4] for it in items)
    amts = "\n".join(it[5] for it in items)

    cols = [c.strip() for c in SP_TABLE_HEADER.split("|") if c.strip()]
    md_lines = [
        "| " + " | ".join(cols) + " |",
        "|" + " --- |" * len(cols),
    ]
    row = (
        f"| {company} | {date} | {inv} | {sales} | {pay_terms} | {due} | "
        f"{codes} | {descs} | {qtys} | {units} | {discs} | {amts} | "
        f"{total} | {vat} | {net_total} |"
    )
    md_lines.append(row)
    return "\n".join(md_lines)


def run_ocr_pipeline(
    *,
    uploaded_bytes: bytes,
    filename: str = "",
    pdf_password: Optional[str] = None,
    api_key: str = "",
    model: str = "",
    task_type: str = "",
    max_tokens: int = 16384,
    temperature: float = 0.1,
    top_p: float = 0.6,
    repetition_penalty: float = 1.2,
    pages_raw: str = "",
    progress_callback: Optional[Callable[[int, int, int], None]] = None,
    page_done_callback: Optional[Callable[[int, int, int, float], None]] = None,
) -> dict[str, Any]:
    temp_path = ""
    start_time = time.perf_counter()
    try:
        temp_path, file_type = save_uploaded_file(
            uploaded_bytes,
            filename or "document.pdf",
            pdf_password,
        )

        if file_type == "image":
            pages_value = [1]
            use_native = False
        else:
            if pages_raw:
                pages_value = parse_pages_input(pages_raw)
                if pages_value is None:
                    page_count = get_pdf_page_count(temp_path)
                    pages_value = list(range(1, page_count + 1))
            else:
                page_count = get_pdf_page_count(temp_path)
                pages_value = list(range(1, page_count + 1))
            use_native = True

        extracted_text, page_texts, page_timings = call_typhoon_ocr(
            file_path=temp_path,
            api_key=api_key,
            model=model,
            task_type=task_type,
            max_tokens=max_tokens,
            temperature=temperature,
            top_p=top_p,
            repetition_penalty=repetition_penalty,
            pages=pages_value,
            progress_callback=progress_callback,
            page_done_callback=page_done_callback,
            use_native_extraction=use_native,
        )

        # แก้คำที่ OCR อ่านผิดในรายละเอียด (เช่น ล้างอดีต→ล้างอัดฉีด, ฤดูฝน→ดูดฝุ่น)
        extracted_text = correct_ocr_description_text(extracted_text)
        page_texts = [correct_ocr_description_text(pt) for pt in page_texts]
        # เว้นบรรทัดก่อนบล็อกสรุปในข้อความที่โชว์ (รวมเงิน, VAT, รวมสุทธิ)
        extracted_text = ensure_summary_spacing_in_text(extracted_text)
        page_texts = [ensure_summary_spacing_in_text(pt) for pt in page_texts]

        extracted_html = render_ocr_html(extracted_text)
        page_htmls = [render_ocr_html(pt) for pt in page_texts]
        extracted_text_b64 = base64.b64encode(extracted_text.encode("utf-8")).decode("utf-8")
        extracted_html_b64 = base64.b64encode(extracted_html.encode("utf-8")).decode("utf-8")
        page_htmls_b64 = base64.b64encode(
            json.dumps(page_htmls, ensure_ascii=False).encode("utf-8")
        ).decode("utf-8")
        elapsed_seconds = round(time.perf_counter() - start_time, 2)

        return {
            "extracted_text": extracted_text,
            "extracted_html": extracted_html,
            "extracted_text_b64": extracted_text_b64,
            "extracted_html_b64": extracted_html_b64,
            "page_htmls_b64": page_htmls_b64,
            "page_texts": page_texts,
            "page_htmls": page_htmls,
            "page_timings": page_timings,
            "elapsed_seconds": elapsed_seconds,
        }
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)


def run_ocr_job(job_id: str, params: dict[str, Any]) -> None:
    try:
        def on_progress(current_step: int, total_steps: int, page_number: int) -> None:
            update_ocr_job(
                job_id,
                status="running",
                message=f"กำลัง OCR หน้า {current_step}/{total_steps}",
                current_step=current_step,
                total_steps=total_steps,
                current_page_number=page_number,
            )

        def on_page_done(
            current_step: int,
            total_steps: int,
            page_number: int,
            elapsed_seconds: float,
        ) -> None:
            append_ocr_job_page_timing(job_id, page_number, elapsed_seconds)
            update_ocr_job(
                job_id,
                status="running",
                message=f"เสร็จหน้า {current_step}/{total_steps} (หน้าเอกสาร {page_number})",
                current_step=current_step,
                total_steps=total_steps,
                current_page_number=page_number,
            )

        result = run_ocr_pipeline(
            uploaded_bytes=params["uploaded_bytes"],
            filename=params.get("filename", ""),
            pdf_password=params.get("pdf_password"),
            api_key=params["api_key"],
            model=params["model"],
            task_type=params["task_type"],
            max_tokens=params["max_tokens"],
            temperature=params["temperature"],
            top_p=params["top_p"],
            repetition_penalty=params["repetition_penalty"],
            pages_raw=params.get("pages_raw", ""),
            progress_callback=on_progress,
            page_done_callback=on_page_done,
        )
        result_id = store_ocr_result(result)
        update_ocr_job(
            job_id,
            status="completed",
            message="OCR เสร็จแล้ว",
            result=result,
            result_id=result_id,
        )
    except Exception as exc:
        update_ocr_job(
            job_id,
            status="failed",
            message="OCR ล้มเหลว",
            error=str(exc),
        )


@app.route("/ocr/start", methods=["POST"])
def ocr_start():
    uploaded_file = request.files.get("pdf_file")
    if not uploaded_file or uploaded_file.filename == "":
        return jsonify({"ok": False, "error": "กรุณาเลือกไฟล์ PDF หรือรูปสแกนก่อน"}), 400

    api_key = request.form.get("api_key", "").strip() or os.getenv("TYPHOON_API_KEY", "")
    if not api_key:
        return jsonify({"ok": False, "error": "กรุณาใส่ Typhoon API Key"}), 400

    params = {
        "uploaded_bytes": uploaded_file.read(),
        "filename": uploaded_file.filename or "",
        "pdf_password": request.form.get("pdf_password", "").strip(),
        "api_key": api_key,
        "model": request.form.get("model", "typhoon-ocr").strip(),
        "task_type": request.form.get("task_type", "default").strip(),
        "max_tokens": int(request.form.get("max_tokens", "16384")),
        "temperature": float(request.form.get("temperature", "0.1")),
        "top_p": float(request.form.get("top_p", "0.6")),
        "repetition_penalty": float(request.form.get("repetition_penalty", "1.2")),
        "pages_raw": request.form.get("pages", "").strip(),
    }

    job_id = uuid.uuid4().hex
    init_ocr_job(job_id)
    worker = threading.Thread(target=run_ocr_job, args=(job_id, params), daemon=True)
    worker.start()
    return jsonify({"ok": True, "job_id": job_id})


@app.route("/ocr/status/<job_id>", methods=["GET"])
def ocr_status(job_id: str):
    with OCR_JOBS_LOCK:
        job = OCR_JOBS.get(job_id)
        if not job:
            return jsonify({"ok": False, "error": "ไม่พบงาน OCR"}), 404
        response = {
            "ok": True,
            "status": job["status"],
            "message": job.get("message", ""),
            "current_step": job.get("current_step", 0),
            "total_steps": job.get("total_steps", 0),
            "current_page_number": job.get("current_page_number", 0),
            "page_timings": job.get("page_timings", []),
            "result_id": job.get("result_id", ""),
            "error": job.get("error", ""),
        }
        if job["status"] == "completed" and job.get("result") is not None:
            response["result"] = job["result"]
    return jsonify(response)


def check_doc_no_exists_in_db(doc_no: str, table_base: str, db_name: str) -> bool:
    """ตรวจสอบว่า เลขที่ (doc_no) มีในตาราง header แล้วหรือยัง"""
    if not doc_no or not table_base:
        return False
    try:
        import pyodbc
    except ModuleNotFoundError:
        return False
    safe_db = sanitize_db_name(db_name)
    safe_base = sanitize_table_name(table_base)
    t_header = f"{safe_base}_header"
    _, target_conn_str = get_sql_connection_strings(safe_db)
    try:
        conn = pyodbc.connect(target_conn_str)
        cur = conn.cursor()
        cur.execute(
            f"SELECT 1 FROM dbo.[{t_header}] WHERE [เลขที่] = ?;",
            (doc_no,),
        )
        row = cur.fetchone()
        conn.close()
        return row is not None
    except pyodbc.Error:
        return False


@app.route("/upload/db/check", methods=["POST"])
def upload_db_check():
    """ตรวจสอบว่าเลขที่จากผล OCR ซ้ำกับข้อมูลในฐานข้อมูลหรือไม่"""
    result_id = request.form.get("result_id", "").strip()
    table_name = request.form.get("table_name", "").strip()
    db_name = request.form.get("db_name", "").strip()
    if not result_id:
        return jsonify({"ok": False, "error": "ไม่พบผลลัพธ์ OCR"}), 400

    result = get_ocr_result(result_id)
    if not result:
        return jsonify({"ok": False, "error": "ผลลัพธ์หมดอายุหรือไม่พบ"}), 404

    try:
        data = extract_ocr_tables_structured(
            result.get("extracted_html", ""),
            result.get("extracted_text", "") or "",
            result.get("page_htmls", []) or [],
        )
        doc_no = (data.get("header") or {}).get("เลขที่", "").strip()
        exists = check_doc_no_exists_in_db(
            doc_no,
            table_name or "OCR_TTB_WEB",
            db_name or "ExcelTtbDB",
        )
        return jsonify({"ok": True, "exists": exists, "doc_no": doc_no or ""})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


@app.route("/upload/db", methods=["POST"])
def upload_db():
    result_id = request.form.get("result_id", "").strip()
    table_name = request.form.get("table_name", "").strip()
    db_name = request.form.get("db_name", "").strip()
    edited_sheets_json = request.form.get("edited_sheets_json", "").strip()
    if not result_id:
        return jsonify({"ok": False, "error": "ไม่พบผลลัพธ์ OCR สำหรับอัพโหลด"}), 400

    result = get_ocr_result(result_id)
    if not result:
        return jsonify({"ok": False, "error": "ผลลัพธ์หมดอายุหรือไม่พบในหน่วยความจำ"}), 404

    try:
        upload_info = upload_result_to_sql_server(
            result, table_name, db_name=db_name or "ExcelTtbDB", edited_sheets_json=edited_sheets_json or None
        )
        return jsonify({"ok": True, "upload_info": upload_info})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


@app.route("/", methods=["GET", "POST"])
def index():
    extracted_text = ""
    extracted_html = ""
    extracted_text_b64 = ""
    extracted_html_b64 = ""
    page_htmls_b64 = ""
    page_texts = []
    page_htmls = []
    page_timings = []
    result_id = ""
    elapsed_seconds = None
    error = ""

    defaults = {
        "model": "typhoon-ocr",
        "task_type": "default",
        "max_tokens": "16384",
        "temperature": "0.1",
        "top_p": "0.6",
        "repetition_penalty": "1.2",
    }

    if request.method == "POST":
        uploaded_file = request.files.get("pdf_file")
        pdf_password = request.form.get("pdf_password", "").strip()
        api_key = request.form.get("api_key", "").strip() or os.getenv("TYPHOON_API_KEY", "")

        model = request.form.get("model", defaults["model"]).strip()
        task_type = request.form.get("task_type", defaults["task_type"]).strip()
        max_tokens = int(request.form.get("max_tokens", defaults["max_tokens"]))
        temperature = float(request.form.get("temperature", defaults["temperature"]))
        top_p = float(request.form.get("top_p", defaults["top_p"]))
        repetition_penalty = float(
            request.form.get("repetition_penalty", defaults["repetition_penalty"])
        )
        pages_raw = request.form.get("pages", "").strip()

        if not uploaded_file or uploaded_file.filename == "":
            error = "กรุณาเลือกไฟล์ PDF หรือรูปสแกนก่อน"
        elif not api_key:
            error = "กรุณาใส่ Typhoon API Key"
        else:
            try:
                result = run_ocr_pipeline(
                    uploaded_bytes=uploaded_file.read(),
                    filename=uploaded_file.filename or "",
                    pdf_password=pdf_password or None,
                    api_key=api_key,
                    model=model,
                    task_type=task_type,
                    max_tokens=max_tokens,
                    temperature=temperature,
                    top_p=top_p,
                    repetition_penalty=repetition_penalty,
                    pages_raw=pages_raw,
                )
                extracted_text = result["extracted_text"]
                extracted_html = result["extracted_html"]
                extracted_text_b64 = result["extracted_text_b64"]
                extracted_html_b64 = result["extracted_html_b64"]
                page_htmls_b64 = result["page_htmls_b64"]
                page_texts = result["page_texts"]
                page_htmls = result["page_htmls"]
                page_timings = result["page_timings"]
                elapsed_seconds = result["elapsed_seconds"]
                result_id = store_ocr_result(result)
            except Exception as exc:  # keep UI simple
                error = str(exc)

    return render_template(
        "index.html",
        extracted_text=extracted_text,
        extracted_html=extracted_html,
        extracted_text_b64=extracted_text_b64,
        extracted_html_b64=extracted_html_b64,
        page_htmls_b64=page_htmls_b64,
        page_texts=page_texts,
        page_htmls=page_htmls,
        page_timings=page_timings,
        result_id=result_id,
        elapsed_seconds=elapsed_seconds,
        error=error,
        defaults=defaults,
    )


@app.route("/download/word", methods=["POST"])
def download_word():
    result_id = request.form.get("result_id", "").strip()
    cached_result = get_ocr_result(result_id)
    if cached_result:
        extracted_html = cached_result.get("extracted_html", "")
        extracted_text = cached_result.get("extracted_text", "")
        page_htmls = cached_result.get("page_htmls", [])
    else:
        extracted_html = decode_base64_payload(request.form.get("extracted_html_b64", ""))
        extracted_text = decode_base64_payload(request.form.get("extracted_text_b64", ""))
        page_htmls = decode_base64_json_list(request.form.get("page_htmls_b64", ""))
    file_data = export_tables_to_docx(extracted_html, extracted_text, page_htmls=page_htmls)
    return send_file(
        file_data,
        as_attachment=True,
        download_name="ocr-tables.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _excel_sheets_to_json(buffer: io.BytesIO) -> list[dict]:
    """Read Excel from buffer and return list of { name, rows } for preview."""
    buffer.seek(0)
    wb = load_workbook(buffer, read_only=True, data_only=True)
    sheets = []
    try:
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append([str(c) if c is not None else "" for c in row])
            sheets.append({"name": sheet.title, "rows": rows})
    finally:
        wb.close()
    return sheets


@app.route("/preview/excel", methods=["POST"])
def preview_excel():
    result_id = request.form.get("result_id", "").strip()
    cached_result = get_ocr_result(result_id)
    if cached_result:
        extracted_html = cached_result.get("extracted_html", "")
        extracted_text = cached_result.get("extracted_text", "")
        page_htmls = cached_result.get("page_htmls", [])
    else:
        extracted_html = decode_base64_payload(request.form.get("extracted_html_b64", ""))
        extracted_text = decode_base64_payload(request.form.get("extracted_text_b64", ""))
        page_htmls = decode_base64_json_list(request.form.get("page_htmls_b64", ""))
    buffer = export_tables_to_excel(extracted_html, extracted_text, page_htmls=page_htmls)
    sheets = _excel_sheets_to_json(buffer)
    return jsonify({"sheets": sheets})


def _excel_from_sheets_json(sheets_json: str) -> io.BytesIO:
    """Build Excel file from JSON payload (list of { name, rows })."""
    data = json.loads(sheets_json or "{}")
    sheets_payload = data.get("sheets") or []
    wb = Workbook()
    default_name = "OCR Tables"
    for idx, sh in enumerate(sheets_payload):
        name = (sh.get("name") or default_name).strip() or default_name
        rows = sh.get("rows") or []
        if idx == 0:
            ws = wb.active
            ws.title = name[:31]
        else:
            ws = wb.create_sheet(title=name[:31])
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, value in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=value)
                if r_idx == 1 and idx == 0:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="D9E2EC", end_color="D9E2EC", fill_type="solid")
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


@app.route("/download/excel", methods=["POST"])
def download_excel():
    sheets_json = request.form.get("sheets_json", "").strip()
    if sheets_json:
        try:
            file_data = _excel_from_sheets_json(sheets_json)
            return send_file(
                file_data,
                as_attachment=True,
                download_name="ocr-tables.xlsx",
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except (json.JSONDecodeError, KeyError, TypeError) as e:
            return jsonify({"error": f"ข้อมูลที่แก้ไขไม่ถูกต้อง: {e!s}"}), 400
    result_id = request.form.get("result_id", "").strip()
    cached_result = get_ocr_result(result_id)
    if cached_result:
        extracted_html = cached_result.get("extracted_html", "")
        extracted_text = cached_result.get("extracted_text", "")
        page_htmls = cached_result.get("page_htmls", [])
    else:
        extracted_html = decode_base64_payload(request.form.get("extracted_html_b64", ""))
        extracted_text = decode_base64_payload(request.form.get("extracted_text_b64", ""))
        page_htmls = decode_base64_json_list(request.form.get("page_htmls_b64", ""))
    file_data = export_tables_to_excel(extracted_html, extracted_text, page_htmls=page_htmls)
    return send_file(
        file_data,
        as_attachment=True,
        download_name="ocr-tables.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=5000, debug=True)
